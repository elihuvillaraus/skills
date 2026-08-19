#!/usr/bin/env python3
"""
La misma propuesta, en Word: Markdown editable -> .docx con la marca del proyecto.

    python3 build_docx.py --fuente propuesta.md --config config.json --out-dir salida/

El Markdown es la fuente, igual que en `build_html_pdf.py`. Este script no inventa
contenido: lo viste para Word, que suele ser el formato en el que el cliente lo edita.

Tipografía: por defecto Calibri (o lo que `fonts.docx_font` diga en config.json), a
propósito — las fuentes de display/cuerpo usadas en el HTML viven en la web y Word
las sustituye por lo que encuentre en la máquina de quien abre el archivo; el color
de marca y la estructura sí viajan.

Generalizado a partir del script real usado en TRUTH/Telmex
(docs/sales/telmex/build_propuesta_docx.py, speechlytics) — misma calidad, marca
parametrizada. No usa pandoc: es un conversor Markdown -> OOXML de mano, porque el
original necesitaba control fino (portada de marca, tablas con sombreado por fila,
índice con folio en color de acento) que pandoc no da sin una plantilla .docx aparte.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ───────────────────────── utilidades de OOXML ─────────────────────────


def sombrear(celda, hexcolor: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    celda._tc.get_or_add_tcPr().append(shd)


def bordes_de_tabla(tabla, hexcolor: str) -> None:
    """Sólo reglas horizontales interiores y una inferior. Nada de rejilla."""
    borders = OxmlElement("w:tblBorders")
    for lado, val in (
        ("top", "none"),
        ("left", "none"),
        ("bottom", "single"),
        ("right", "none"),
        ("insideH", "single"),
        ("insideV", "none"),
    ):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hexcolor)
        borders.append(el)
    tabla._tbl.tblPr.append(borders)


def fijar_anchos(tabla, anchos_cm: list[float]) -> None:
    """python-docx sólo respeta el ancho si se pone celda por celda y sin autofit."""
    tabla.autofit = False
    for fila in tabla.rows:
        for celda, cm in zip(fila.cells, anchos_cm):
            celda.width = Cm(cm)


def repetir_encabezado(fila) -> None:
    trPr = fila._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def espaciado_de_letra(run, twentieths: int) -> None:
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(twentieths))
    run._element.get_or_add_rPr().append(el)


def sin_separar(parrafo) -> None:
    """`keep with next`: un encabezado no se queda solo al pie de la página."""
    parrafo.paragraph_format.keep_with_next = True


def hex_a_rgb(hexcolor: str) -> RGBColor:
    h = hexcolor.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ───────────────────────── inline: **negrita**, `código` ─────────────────────────

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def escribir_inline(parrafo, texto: str, base_size: Pt, color: RGBColor, paleta: dict, fuente_txt: str) -> None:
    for trozo in INLINE.split(texto):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            r = parrafo.add_run(trozo[2:-2])
            r.bold = True
            r.font.color.rgb = paleta["tinta"]
        elif trozo.startswith("`") and trozo.endswith("`"):
            r = parrafo.add_run(trozo[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = color
        elif trozo.startswith("*") and trozo.endswith("*") and len(trozo) > 2:
            r = parrafo.add_run(trozo[1:-1])
            r.italic = True
            r.font.color.rgb = paleta["tenue"]
        else:
            r = parrafo.add_run(trozo)
            r.font.color.rgb = color
        r.font.size = base_size
        r.font.name = r.font.name or fuente_txt


# ───────────────────────── troceado del Markdown ─────────────────────────


def bloques(md: str):
    """Devuelve (tipo, contenido) en orden. Tipos: h1..h4, p, ul, ol, tabla, hr."""
    lineas = md.split("\n")
    i, n = 0, len(lineas)
    while i < n:
        crudo = lineas[i].rstrip()

        if not crudo.strip():
            i += 1
            continue

        if crudo == "---":
            yield ("hr", None)
            i += 1
            continue

        m = re.match(r"^(#{1,6}) (.+)$", crudo)
        if m:
            yield (f"h{min(len(m.group(1)), 4)}", m.group(2).strip())
            i += 1
            continue

        if crudo.startswith("|"):
            filas = []
            while i < n and lineas[i].startswith("|"):
                filas.append(lineas[i])
                i += 1
            yield ("tabla", filas)
            continue

        if re.match(r"^[-*] ", crudo) or re.match(r"^\d+\. ", crudo):
            ordenada = bool(re.match(r"^\d+\. ", crudo))
            items = []
            while i < n:
                c = lineas[i].rstrip()
                if re.match(r"^[-*] ", c) or re.match(r"^\d+\. ", c):
                    items.append(re.sub(r"^(?:[-*] |\d+\. )", "", c))
                    i += 1
                elif c.startswith("  ") and c.strip() and items:
                    items[-1] += " " + c.strip()
                    i += 1
                else:
                    break
            yield ("ol" if ordenada else "ul", items)
            continue

        partes = []
        while i < n:
            c = lineas[i].rstrip()
            if (
                not c.strip()
                or c == "---"
                or c.startswith("|")
                or c.startswith("#")
                or re.match(r"^[-*] ", c)
                or re.match(r"^\d+\. ", c)
            ):
                break
            partes.append(c.strip())
            i += 1
        if not partes:
            i += 1
            continue
        yield ("p", " ".join(partes))


def parsear_tabla(filas: list[str]):
    def celdas(l: str):
        return [c.strip() for c in l.strip().strip("|").split("|")]

    encabezado = celdas(filas[0])
    if len(filas) > 1 and re.match(r"^\|[\s:|-]+\|$", filas[1].strip()):
        alineacion = ["right" if spec.endswith(":") and not spec.startswith(":") else "left" for spec in celdas(filas[1])]
        cuerpo = [celdas(l) for l in filas[2:]]
    else:
        alineacion = ["left"] * len(encabezado)
        cuerpo = [celdas(l) for l in filas[1:]]
    return encabezado, alineacion, cuerpo


# ───────────────────────── construcción del documento ─────────────────────────


def paleta_desde_config(cfg: dict) -> dict:
    c = cfg["colors"]
    return {
        "primary": hex_a_rgb(c["primary"]),
        "accent": hex_a_rgb(c["accent"]),
        "tinta": hex_a_rgb(c["text_paper"]),
        "cuerpo": hex_a_rgb(c["text"]),
        "tenue": hex_a_rgb(c["text_dim"]),
        "gris": hex_a_rgb(c["muted"]),
        "primary_hex": c["primary"].lstrip("#").upper(),
        "accent_hex": c["accent"].lstrip("#").upper(),
        "fila_hex": c["surface"].lstrip("#").upper(),
        "regla_hex": c["hairline"].lstrip("#").upper(),
    }


def preparar(doc: Document, fuente_txt: str, paleta: dict) -> None:
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = fuente_txt
    normal.font.size = Pt(10)
    normal.font.color.rgb = paleta["cuerpo"]
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15


def regla_acento(doc: Document, paleta: dict, ancho_cm: float = 2.2) -> None:
    t = doc.add_table(rows=1, cols=1)
    fijar_anchos(t, [ancho_cm])
    c = t.cell(0, 0)
    sombrear(c, paleta["accent_hex"])
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run("").font.size = Pt(2)


def portada(doc: Document, cfg: dict, root: Path, paleta: dict) -> None:
    brand = cfg["brand"]
    cover = cfg["cover"]

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(brand["name"])
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = paleta["primary"]
    espaciado_de_letra(r, 70)
    if brand.get("name_accent_suffix"):
        r2 = p.add_run(brand["name_accent_suffix"])
        r2.bold = True
        r2.font.size = Pt(16)
        r2.font.color.rgb = paleta["accent"]

    regla_acento(doc, paleta)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(cover.get("kicker_docx", cover["kicker"]).upper())
    r.font.size = Pt(8.5)
    r.font.color.rgb = paleta["tenue"]
    espaciado_de_letra(r, 60)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    titulo_plano = re.sub(r"<[^>]+>", "", cover["title_html"])
    titulo_plano = re.sub(r"&iacute;", "í", titulo_plano)
    titulo_plano = re.sub(r"&eacute;", "é", titulo_plano)
    titulo_plano = re.sub(r"&aacute;", "á", titulo_plano)
    titulo_plano = re.sub(r"&oacute;", "ó", titulo_plano)
    titulo_plano = re.sub(r"&uacute;", "ú", titulo_plano)
    negrita = cover.get("title_bold_docx", "")
    if negrita and negrita in titulo_plano:
        antes, despues = titulo_plano.split(negrita, 1)
        r = p.add_run(antes)
        r.font.size = Pt(26)
        r.font.color.rgb = paleta["tinta"]
        r = p.add_run(negrita)
        r.font.size = Pt(26)
        r.bold = True
        r.font.color.rgb = paleta["primary"]
        r = p.add_run(despues)
        r.font.size = Pt(26)
        r.font.color.rgb = paleta["tinta"]
    else:
        r = p.add_run(titulo_plano)
        r.font.size = Pt(26)
        r.font.color.rgb = paleta["tinta"]

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(cover["subtitle"])
    r.font.size = Pt(11.5)
    r.font.color.rgb = paleta["tenue"]

    for _ in range(6):
        doc.add_paragraph()

    meta = cover.get("meta", [])
    if meta:
        t = doc.add_table(rows=2, cols=len(meta))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        fijar_anchos(t, [16.6 / len(meta)] * len(meta))
        for j, (rotulo, valor) in enumerate(meta):
            p = t.cell(0, j).paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(rotulo.upper())
            r.font.size = Pt(7)
            r.font.color.rgb = paleta["gris"]
            espaciado_de_letra(r, 30)
            p = t.cell(1, j).paragraphs[0]
            r = p.add_run(valor)
            r.font.size = Pt(9)
            r.bold = True
            r.font.color.rgb = paleta["tinta"]

    doc.add_paragraph()

    logo_path = root / brand["logo_path"] if brand.get("logo_path") else None
    firma = doc.add_table(rows=1, cols=2)
    fijar_anchos(firma, [4.8, 11.8])
    if logo_path and logo_path.exists():
        firma.cell(0, 0).paragraphs[0].add_run().add_picture(str(logo_path), width=Cm(4.2))
    p = firma.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    tagline_bold = brand.get("tagline_bold", "")
    tagline = brand.get("tagline", "")
    if tagline_bold and tagline_bold in tagline:
        antes, despues = tagline.split(tagline_bold, 1)
        for texto, negrita_ in ((antes, False), (tagline_bold, True), (despues, False)):
            if not texto:
                continue
            r = p.add_run(texto)
            r.bold = negrita_
            r.font.size = Pt(8.5)
            r.font.color.rgb = paleta["tinta"] if negrita_ else paleta["tenue"]
    else:
        r = p.add_run(tagline)
        r.font.size = Pt(8.5)
        r.font.color.rgb = paleta["tenue"]

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def indice(doc: Document, md: str, paleta: dict) -> None:
    p = doc.add_paragraph()
    sin_separar(p)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Contenido")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = paleta["tinta"]

    titulos = re.findall(r"^## (.+)$", md, flags=re.MULTILINE)
    t = doc.add_table(rows=0, cols=2)
    bordes_de_tabla(t, paleta["regla_hex"])
    for titulo in titulos:
        titulo = titulo.strip()
        fila = t.add_row()
        m = re.match(r"^(\d+)\.\s+(.*)$", titulo)
        pn = fila.cells[0].paragraphs[0]
        pn.paragraph_format.space_before = Pt(4)
        pn.paragraph_format.space_after = Pt(4)
        pt = fila.cells[1].paragraphs[0]
        pt.paragraph_format.space_before = Pt(4)
        pt.paragraph_format.space_after = Pt(4)
        if m:
            r = pn.add_run(m.group(1))
            r.font.size = Pt(9)
            r.font.color.rgb = paleta["accent"]
            r = pt.add_run(m.group(2))
            r.font.size = Pt(10.5)
            r.font.color.rgb = paleta["tinta"]
        else:
            r = pt.add_run(titulo)
            r.font.size = Pt(10)
            r.font.color.rgb = paleta["tenue"]

    fijar_anchos(t, [1.4, 15.2])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def escribir_tabla(doc: Document, filas: list[str], paleta: dict, fuente_txt: str) -> None:
    encabezado, alineacion, cuerpo = parsear_tabla(filas)
    t = doc.add_table(rows=1, cols=len(encabezado))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    bordes_de_tabla(t, paleta["regla_hex"])

    for j, celda in enumerate(encabezado):
        c = t.rows[0].cells[j]
        sombrear(c, paleta["primary_hex"])
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if alineacion[j] == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(re.sub(r"\*\*", "", celda))
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    repetir_encabezado(t.rows[0])

    for k, fila in enumerate(cuerpo):
        r_ = t.add_row()
        for j in range(len(encabezado)):
            texto = fila[j] if j < len(fila) else ""
            c = r_.cells[j]
            if k % 2 == 1:
                sombrear(c, paleta["fila_hex"])
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if alineacion[j] == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            escribir_inline(p, texto, Pt(9), paleta["cuerpo"], paleta, fuente_txt)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def primer_encabezado_numerado(md: str) -> int:
    m = re.search(r"^## \d+\.\s", md, flags=re.MULTILINE)
    return m.start() if m else -1


def construir(md_path: Path, cfg: dict, root: Path, salida: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    i = primer_encabezado_numerado(md)
    cuerpo_md = md if i == -1 else md[i:]

    fuente_txt = cfg["fonts"].get("docx_font", "Calibri")
    paleta = paleta_desde_config(cfg)

    doc = Document()
    preparar(doc, fuente_txt, paleta)
    portada(doc, cfg, root, paleta)
    indice(doc, cuerpo_md, paleta)

    primera_h2 = True
    for tipo, contenido in bloques(cuerpo_md):
        if tipo == "h2":
            if not primera_h2:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            primera_h2 = False
            p = doc.add_paragraph()
            sin_separar(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(10)
            pbdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "12")
            top.set(qn("w:space"), "8")
            top.set(qn("w:color"), paleta["primary_hex"])
            pbdr.append(top)
            p._p.get_or_add_pPr().append(pbdr)
            r = p.add_run(contenido)
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = paleta["tinta"]

        elif tipo == "h3":
            p = doc.add_paragraph()
            sin_separar(p)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(contenido)
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = paleta["primary"]

        elif tipo == "h4":
            p = doc.add_paragraph()
            sin_separar(p)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(contenido)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = paleta["tinta"]

        elif tipo == "h1":
            p = doc.add_paragraph()
            sin_separar(p)
            r = p.add_run(contenido)
            r.bold = True
            r.font.size = Pt(18)
            r.font.color.rgb = paleta["primary"]

        elif tipo == "p":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            escribir_inline(p, contenido, Pt(10), paleta["cuerpo"], paleta, fuente_txt)

        elif tipo in ("ul", "ol"):
            for item in contenido:
                p = doc.add_paragraph(style="List Number" if tipo == "ol" else "List Bullet")
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.7)
                escribir_inline(p, item, Pt(10), paleta["cuerpo"], paleta, fuente_txt)

        elif tipo == "tabla":
            escribir_tabla(doc, contenido, paleta, fuente_txt)

        elif tipo == "hr":
            pass

    doc.save(str(salida))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--fuente", type=Path, required=True, help="Markdown fuente de la propuesta")
    ap.add_argument("--config", type=Path, required=True, help="config.json con marca/colores/portada")
    ap.add_argument("--out-dir", type=Path, required=True, help="carpeta de salida")
    args = ap.parse_args()

    if not args.fuente.exists():
        print(f"No existe {args.fuente}", file=sys.stderr)
        return 1
    cfg = json.loads(args.config.read_text(encoding="utf-8"))
    root = args.config.resolve().parent
    args.out_dir.mkdir(parents=True, exist_ok=True)

    salida = args.out_dir / f"{cfg['output']['base_name']}.docx"
    construir(args.fuente, cfg, root, salida)
    print(f"  DOCX  {salida.name}  ({salida.stat().st_size / 1024:.0f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
